import fitz  # PyMuPDF

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI

from langchain_chroma import Chroma 

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import os
import io

from dotenv import load_dotenv

import json

from docx import Document as DocxDocument # to avoid conflict with document import from langchain

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
USER_CREDENTIALS_FILE = 'config/credentials.json'
TOKEN_PATH = 'token.json'

def authenticate_drive():
    creds = None

    # 1. Load existing token if available
    if os.path.exists(TOKEN_PATH):
        creds = Credentials.from_authorized_user_file(TOKEN_PATH, SCOPES)

    # 2. Refresh token if expired and refresh_token is available
    if creds and creds.expired and creds.refresh_token:
        creds.refresh(Request())

    # 3. Do full browser login if no creds or can't refresh
    if not creds or not creds.valid:
        flow = InstalledAppFlow.from_client_secrets_file(USER_CREDENTIALS_FILE, SCOPES)
        creds = flow.run_local_server(port=0)

        # Save token for future use
        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())

    return build('drive', 'v3', credentials=creds)


def list_files(service, mime_type_filter=None):
    print("listing files")
    query = f"mimeType='{mime_type_filter}'" if mime_type_filter else ""
    results = service.files().list(q=query, pageSize=10, fields="files(id, name, mimeType)").execute()
    return results.get('files', [])


def download_file(service, file_id, file_name):
    request = service.files().get_media(fileId=file_id)
    fh = io.FileIO(file_name, 'wb')
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()


def extract_text_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_text_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_text_txt(file_path):
    with open(file_path, 'r') as file:
        return file.read()


def extracted_text(f, service):
    # print(f['id'], "  ", f['name'])
    
    print(f"Processing: {f['name']}")
    download_file(service, f['id'], f['name'])
    # print("file downloaded")
    if f['name'].endswith('.pdf'):
        return extract_text_pdf(f['name'])
    elif f['name'].endswith('.docx'):
        return extract_text_docx(f['name'])
    elif f['name'].endswith('.txt'):
        return extract_text_txt(f['name'])



PROCESSED_IDS_FILE = "config/processed_gd_files.json"

def load_gd_processed_ids():
    if not os.path.exists(PROCESSED_IDS_FILE):
        return set()
    with open(PROCESSED_IDS_FILE, "r") as file:
        try:
            return set(json.load(file))
        except json.JSONDecodeError:
            return set()

def save_gd_processed_id(file_id):
    processed_ids = load_gd_processed_ids()
    processed_ids.add(file_id)
    with open(PROCESSED_IDS_FILE, "w") as file:
        json.dump(list(processed_ids), file)



def split_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
    return splitter.split_text(text)

def create_documents(chunks, metadata):
    return [Document(page_content=chunk, metadata=metadata) for chunk in chunks]


load_dotenv()

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    raise ValueError("❌ GOOGLE_API_KEY not found. Make sure it's set in .env.")


embeddings = GoogleGenerativeAIEmbeddings(
    model="models/embedding-001",
    google_api_key=GOOGLE_API_KEY
)

CHROMA_DIR = "vector_store/chroma"
from uuid import uuid4

def chunk_and_store(text, metadata, embeddings):
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
    chunks = splitter.split_text(text)

    documents = []
    for i, chunk in enumerate(chunks):
        chunk_meta = metadata.copy()
        chunk_meta.update({
            "chunk_index": i,
            "chunk_id": str(uuid4()),  # useful for deduplication or updates
        })
        documents.append(Document(page_content=chunk, metadata=chunk_meta))

    if not os.path.exists(CHROMA_DIR):
        db = Chroma.from_documents(documents, embeddings, persist_directory=CHROMA_DIR)
    else:
        db = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)
        db.add_documents(documents)


def process_drive():
    service = authenticate_drive()
    files = list_files(service)
    processed_ids = load_gd_processed_ids()
    for file in files:
        file_id =  file['id']
        file_name = file['name']
        if file_id in processed_ids:
            print(f"⏩ Skipping already processed file : {file_name} {file_id}")
            continue
        

        if file['name'].endswith('.pdf') or file['name'].endswith('.docx') or file['name'].endswith('.txt'):
            try:
                text = extracted_text(file, service)
                print("Content extracted!")
                if not text.strip(): continue
                metadata = {
                    "source": "google_drive",
                    "file_name": file["name"],
                    "file_id": file["id"]
                }
                chunk_and_store(text, metadata, embeddings)
                save_gd_processed_id(file_id)
            except Exception as e:
                print(f"[!] Failed on {file['name']}: {e}")


llm = ChatGoogleGenerativeAI(model="models/gemini-2.0-flash", google_api_key=GOOGLE_API_KEY)


if __name__ == "__main__":
<<<<<<< HEAD
    process_drive()

    # from langchain.retrievers import EnsembleRetriever
    
    # from langchain.chains import RetrievalQA

    # # Two separate Chroma DBs
    # db1 = Chroma(persist_directory="vector_store/chroma", embedding_function=embeddings)
    # db2 = Chroma(persist_directory="vector_store/emails", embedding_function=embeddings)

    # retriever1 = db1.as_retriever(search_kwargs={"k": 3})
    # retriever2 = db2.as_retriever(search_kwargs={"k": 3})

    # ensemble_retriever = EnsembleRetriever(
    #     retrievers=[retriever1, retriever2],
    #     weights=[0.7, 0.3]  # Equal weightage or customize based on trust
    # )

    # qa_chain = RetrievalQA.from_chain_type(
    #     llm=llm,
    #     retriever=ensemble_retriever,
    #     chain_type="stuff"
    # )
    
    # query = "tell me about benefits given to employees"
    # response = qa_chain.invoke({"query": query})
    # print(response["result"])
=======
    # process_drive()

    from langchain.retrievers import EnsembleRetriever
    
    from langchain.chains import RetrievalQA

    # Two separate Chroma DBs
    db1 = Chroma(persist_directory="vector_store/chroma", embedding_function=embeddings)
    db2 = Chroma(persist_directory="vector_store/emails", embedding_function=embeddings)

    retriever1 = db1.as_retriever(search_kwargs={"k": 3})
    retriever2 = db2.as_retriever(search_kwargs={"k": 3})

    ensemble_retriever = EnsembleRetriever(
        retrievers=[retriever1, retriever2],
        weights=[0.7, 0.3]  # Equal weightage or customize based on trust
    )

    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=ensemble_retriever,
        chain_type="stuff"
    )
    
    # db = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)

    # qa_chain = RetrievalQA.from_chain_type(
    #     llm=llm,
    #     chain_type="stuff",  # simplest, can be replaced with map_reduce or refine
    #     retriever=db.as_retriever(search_kwargs={"k": 3})
    # )
    query = "tell me about benefits given to employees"
    response = qa_chain.invoke({"query": query})
    print(response["result"])
>>>>>>> 4c05c94 (emails vector database added)

